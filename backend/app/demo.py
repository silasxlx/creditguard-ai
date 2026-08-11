from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from io import BytesIO
from typing import Any

from sqlalchemy import select
from sqlalchemy.orm import Session

from .models import CreditCase, DocumentType, ReviewRun
from .schemas import CaseCreateRequest, RunCreateRequest
from .service import create_case, create_run, get_case, save_document


@dataclass(frozen=True)
class DemoScenario:
    scenario_id: str
    case_no: str
    customer_name: str
    customer_key: str
    credit_term_months: int
    due_diligence_term_months: int


SCENARIOS: dict[str, DemoScenario] = {
    "DEMO-NORMAL-001": DemoScenario(
        "DEMO-NORMAL-001",
        "DEMO-NORMAL-001",
        "星海演示科技有限公司",
        "SYNTH-DEMO-NORMAL-001",
        24,
        24,
    ),
    "DEMO-HIGH-001": DemoScenario(
        "DEMO-HIGH-001",
        "DEMO-HIGH-001",
        "远山演示制造有限公司",
        "SYNTH-DEMO-HIGH-001",
        24,
        48,
    ),
}


def _fact_lines(scenario: DemoScenario, term_months: int) -> list[str]:
    return [
        f"企业名称：{scenario.customer_name}",
        f"统一社会信用代码：91310000{scenario.scenario_id.replace('-', '')[:10]}",
        "成立日期：2020-01-01",
        "法定代表人：演示负责人",
        "所属行业：制造业",
        "申请金额：1000000",
        "币种：CNY",
        f"申请期限：{term_months}个月",
        "贷款用途：流动资金",
        "总资产：5000000",
        "总负债：2000000",
        "流动资产：3000000",
        "流动负债：1500000",
        "营业收入：8000000",
        "净利润：600000",
    ]


def _pdf_bytes(lines: list[str]) -> bytes:
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_text((48, 48), "\n".join(lines), fontname="china-s", fontsize=9)
    content = document.tobytes()
    document.close()
    return content


def _docx_bytes(lines: list[str]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("CreditGuard AI 合成演示材料", level=1)
    document.add_paragraph("仅用于授信智能合规审查 PoC 演示，不代表真实客户资料。")
    for line in lines:
        document.add_paragraph(line)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes(lines: list[str]) -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    if sheet is None:
        raise RuntimeError("A workbook must have an active worksheet")
    sheet.title = "合成财务报表"
    sheet.append(["字段", "数值"])
    for line in lines:
        key, value = line.split("：", 1)
        sheet.append([key, value])
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _materials(scenario: DemoScenario) -> list[tuple[DocumentType, str, str, bytes]]:
    common = _fact_lines(scenario, scenario.credit_term_months)
    due = _fact_lines(scenario, scenario.due_diligence_term_months)
    return [
        (
            DocumentType.BUSINESS_LICENSE,
            "business-license.pdf",
            "application/pdf",
            _pdf_bytes(common[:5]),
        ),
        (
            DocumentType.CREDIT_APPLICATION,
            "credit-application.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _docx_bytes(common[5:9]),
        ),
        (
            DocumentType.DUE_DILIGENCE,
            "due-diligence.pdf",
            "application/pdf",
            _pdf_bytes([common[0], common[1], common[4], *due[7:10]]),
        ),
        (
            DocumentType.FINANCIAL_STATEMENTS,
            "financial-statements.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            _xlsx_bytes(common[9:]),
        ),
    ]


def scenario_materials(scenario_id: str) -> list[tuple[DocumentType, str, str, bytes]]:
    scenario = SCENARIOS.get(scenario_id)
    if scenario is None:
        raise KeyError(scenario_id)
    return _materials(scenario)


def create_demo_scenario(
    db: Session, scenario_id: str, actor: str, idempotency_key: str
) -> tuple[CreditCase, ReviewRun, list[str], bool]:
    scenario = SCENARIOS.get(scenario_id)
    if scenario is None:
        raise KeyError(scenario_id)

    existing_case = db.scalar(
        select(CreditCase).where(CreditCase.case_no == scenario.case_no)
    )
    if existing_case is not None:
        existing_run = db.scalar(
            select(ReviewRun)
            .where(ReviewRun.case_id == existing_case.id)
            .order_by(ReviewRun.created_at.desc())
        )
        if existing_run is not None:
            return existing_case, existing_run, list(existing_run.document_version_ids), False

    case = create_case(
        db,
        CaseCreateRequest(
            case_no=scenario.case_no,
            customer_name=scenario.customer_name,
            customer_key=scenario.customer_key,
            review_date=date(2026, 8, 11),
        ),
        actor,
        f"{idempotency_key}:case",
    )
    document_ids: list[str] = []
    for index, (document_type, filename, mime, content) in enumerate(_materials(scenario)):
        document = save_document(
            db,
            case_id=case.id,
            document_type=document_type,
            filename=filename,
            mime=mime,
            content=content,
            actor=actor,
            idempotency_key=f"{idempotency_key}:document:{index}",
        )
        document_ids.append(document.id)

    current_case = get_case(db, case.id)
    run = create_run(
        db,
        case.id,
        RunCreateRequest(
            document_version_ids=document_ids,
            expected_case_version=current_case.version,
        ),
        actor,
        f"{idempotency_key}:run",
    )
    return current_case, run, document_ids, True


def demo_response_payload(
    scenario_id: str,
    case: CreditCase,
    run: ReviewRun,
    document_ids: list[str],
    created: bool,
) -> dict[str, Any]:
    return {
        "scenario_id": scenario_id,
        "case_id": case.id,
        "run_id": run.id,
        "case_version": case.version,
        "input_document_version_ids": document_ids,
        "run_status": run.status.value,
        "created": created,
    }
