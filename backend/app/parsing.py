from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Protocol, cast

from .adapters import ParsedBlock


def effective_chars(text: str) -> int:
    return len(re.sub(r"[\s\u0000-\u001f\u007f]", "", text))


def _stable_block_id(document_version_id: str, page: int, reading_order: int, text: str) -> str:
    text_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
    raw = f"{document_version_id}|{page}|{reading_order}|{text_hash}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _evidence_id(document_version_id: str, block_id: str) -> str:
    return hashlib.sha256(f"{document_version_id}|{block_id}".encode()).hexdigest()


class MinerUAdapter(Protocol):
    def parse(
        self, document_version_id: str, content: bytes, filename: str
    ) -> list[ParsedBlock]: ...


@dataclass
class ParseQuality:
    parser: str
    effective_characters: int
    page_count: int
    locatable_characters: int
    average_chars_per_page: float
    needs_mineru: bool = False
    reasons: list[str] = field(default_factory=list)
    rejected: bool = False
    cache_missing: bool = False


@dataclass
class ParsedDocument:
    document_version_id: str
    filename: str
    blocks: list[ParsedBlock]
    quality: ParseQuality


def _make_block(
    document_version_id: str,
    text: str,
    source: str,
    reading_order: int,
    locator: dict[str, object],
    page: int | None = None,
    block_type: str = "text",
    section_path: list[str] | None = None,
    provider: str = "local",
) -> ParsedBlock:
    block_id = _stable_block_id(document_version_id, page or 0, reading_order, text)
    return ParsedBlock(
        document_version_id=document_version_id,
        block_id=block_id,
        evidence_id=_evidence_id(document_version_id, block_id),
        page=page,
        block_type=block_type,
        text=text,
        source=source,
        locator=locator,
        section_path=section_path or [],
        reading_order=reading_order,
        provider=provider,
    )


def parse_pdf(document_version_id: str, content: bytes, filename: str) -> ParsedDocument:
    try:
        import fitz
    except ImportError as exc:  # pragma: no cover - dependency is locked in pyproject
        raise RuntimeError("PyMuPDF is required for PDF parsing") from exc

    document = fitz.open(stream=content, filetype="pdf")
    if document.needs_pass:
        return ParsedDocument(
            document_version_id=document_version_id,
            filename=filename,
            blocks=[],
            quality=ParseQuality(
                "pymupdf", 0, document.page_count, 0, 0.0, rejected=True, reasons=["ENCRYPTED_PDF"]
            ),
        )

    blocks: list[ParsedBlock] = []
    extracted = 0
    locatable = 0
    order = 0
    for page_index in range(document.page_count):
        page = document.load_page(page_index)
        page_number = page_index + 1
        page_dict = cast(dict[str, Any], page.get_text("dict"))
        for raw_block in page_dict.get("blocks", []):
            if raw_block.get("type") != 0:
                continue
            lines = raw_block.get("lines", [])
            text = "\n".join(
                span.get("text", "") for line in lines for span in line.get("spans", [])
            ).strip()
            if not text:
                continue
            chars = effective_chars(text)
            extracted += chars
            locatable += chars
            blocks.append(
                _make_block(
                    document_version_id,
                    text,
                    filename,
                    order,
                    page=page_number,
                    locator={"page": page_number, "bbox": list(raw_block.get("bbox", []))},
                )
            )
            order += 1

    page_count = max(document.page_count, 1)
    average = extracted / page_count
    reasons: list[str] = []
    if extracted == 0:
        reasons.append("NO_TEXT")
    if average < 50:
        reasons.append("LOW_TEXT_DENSITY")
    if extracted and locatable / extracted < 0.90:
        reasons.append("LOW_LOCATABLE_RATIO")
    return ParsedDocument(
        document_version_id=document_version_id,
        filename=filename,
        blocks=blocks,
        quality=ParseQuality(
            "pymupdf",
            extracted,
            document.page_count,
            locatable,
            average,
            needs_mineru=bool(reasons),
            reasons=reasons,
        ),
    )


def parse_docx(document_version_id: str, content: bytes, filename: str) -> ParsedDocument:
    from io import BytesIO

    from docx import Document as DocxDocument

    document = DocxDocument(BytesIO(content))
    blocks: list[ParsedBlock] = []
    order = 0
    section_path: list[str] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = str(paragraph.style.name or "") if paragraph.style is not None else ""
        if style_name.startswith("Heading"):
            level_match = re.search(r"(\d+)$", style_name)
            level = int(level_match.group(1)) if level_match else 1
            section_path = section_path[: max(level - 1, 0)]
            section_path.append(text)
        blocks.append(
            _make_block(
                document_version_id,
                text,
                filename,
                order,
                {"paragraph_index": paragraph_index},
                section_path=list(section_path),
            )
        )
        order += 1
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            values = [cell.text.strip() for cell in row.cells]
            text = " | ".join(value for value in values if value)
            if not text:
                continue
            blocks.append(
                _make_block(
                    document_version_id,
                    text,
                    filename,
                    order,
                    {"table_index": table_index, "row_index": row_index},
                    block_type="table",
                )
            )
            order += 1
    chars = sum(effective_chars(block.text) for block in blocks)
    return ParsedDocument(
        document_version_id,
        filename,
        blocks,
        ParseQuality("python-docx", chars, 1, chars, float(chars)),
    )


def parse_xlsx(document_version_id: str, content: bytes, filename: str) -> ParsedDocument:
    from io import BytesIO

    import openpyxl

    workbook = openpyxl.load_workbook(BytesIO(content), data_only=True, read_only=True)
    formula_workbook = openpyxl.load_workbook(BytesIO(content), data_only=False, read_only=True)
    blocks: list[ParsedBlock] = []
    order = 0
    cache_missing = False
    for sheet, formula_sheet in zip(workbook.worksheets, formula_workbook.worksheets, strict=True):
        for row, formula_row in zip(sheet.iter_rows(), formula_sheet.iter_rows(), strict=True):
            values: list[str] = []
            for cell, formula_cell in zip(row, formula_row, strict=True):
                if (
                    isinstance(formula_cell.value, str)
                    and formula_cell.value.startswith("=")
                    and cell.value is None
                ):
                    cache_missing = True
                if cell.value is None:
                    continue
                values.append(f"{cell.coordinate}: {cell.value}")
            if not values:
                continue
            raw_values = [str(cell.value).strip() for cell in row if cell.value is not None]
            if len(raw_values) >= 2:
                text = f"{raw_values[0]}: {raw_values[1]}"
            else:
                text = " | ".join(values)
            blocks.append(
                _make_block(
                    document_version_id,
                    text,
                    filename,
                    order,
                    {"sheet": sheet.title, "cell_range": row[0].row if row else None},
                    block_type="table",
                )
            )
            order += 1
    chars = sum(effective_chars(block.text) for block in blocks)
    reasons = ["FORMULA_CACHE_MISSING"] if cache_missing else []
    workbook.close()
    formula_workbook.close()
    return ParsedDocument(
        document_version_id,
        filename,
        blocks,
        ParseQuality(
            "openpyxl",
            chars,
            len(workbook.worksheets),
            chars,
            float(chars),
            cache_missing=cache_missing,
            reasons=reasons,
        ),
    )


def parse_mineru_content_list(
    document_version_id: str,
    filename: str,
    content_list: list[dict[str, Any]],
    provider_version: str = "unknown",
) -> ParsedDocument:
    blocks: list[ParsedBlock] = []
    for order, item in enumerate(content_list):
        text = str(item.get("text") or item.get("content") or "").strip()
        if not text:
            continue
        page = int(item.get("page_idx", 0)) + 1
        bbox = item.get("bbox") or []
        blocks.append(
            _make_block(
                document_version_id,
                text,
                filename,
                order,
                {"page": page, "bbox": bbox, "provider_version": provider_version},
                page=page,
                block_type=str(item.get("type") or "text"),
                provider=f"mineru:{provider_version}",
            )
        )
    chars = sum(effective_chars(block.text) for block in blocks)
    pages = max((block.page or 1 for block in blocks), default=1)
    return ParsedDocument(
        document_version_id,
        filename,
        blocks,
        ParseQuality("mineru", chars, pages, chars, chars / pages, reasons=[]),
    )


def _quality_from_blocks(parser: str, blocks: list[ParsedBlock]) -> ParseQuality:
    characters = sum(effective_chars(block.text) for block in blocks)
    pages = max((block.page or 1 for block in blocks), default=1)
    return ParseQuality(parser, characters, pages, characters, characters / pages)


def parse_document(
    document_version_id: str,
    content: bytes,
    filename: str,
    mineru: MinerUAdapter | None = None,
) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".pdf":
            parsed = parse_pdf(document_version_id, content, filename)
            if parsed.quality.needs_mineru:
                if mineru is None:
                    parsed.quality.reasons.append("MINERU_NOT_CONFIGURED")
                    return parsed
                blocks = mineru.parse(document_version_id, content, filename)
                return ParsedDocument(
                    document_version_id,
                    filename,
                    blocks,
                    _quality_from_blocks("mineru", blocks),
                )
            return parsed
        if suffix == ".docx":
            return parse_docx(document_version_id, content, filename)
        if suffix == ".xlsx":
            return parse_xlsx(document_version_id, content, filename)
    except Exception as exc:
        return ParsedDocument(
            document_version_id,
            filename,
            [],
            ParseQuality(
                "unknown",
                0,
                0,
                0,
                0.0,
                needs_mineru=suffix == ".pdf",
                reasons=["PARSE_ERROR", type(exc).__name__],
            ),
        )
    raise ValueError(f"Unsupported document suffix: {suffix}")


class MinerUHttpAdapter:
    """Small boundary adapter; endpoint details stay configurable and auditable."""

    def __init__(self, base_url: str, api_key: str, timeout_seconds: float = 180.0) -> None:
        self.base_url = base_url.rstrip("/")
        self.api_key = api_key
        self.timeout_seconds = timeout_seconds

    def parse(self, document_version_id: str, content: bytes, filename: str) -> list[ParsedBlock]:
        import httpx

        headers = {"Authorization": f"Bearer {self.api_key}"}
        with httpx.Client(timeout=self.timeout_seconds) as client:
            response = client.post(
                f"{self.base_url}/file-parse",
                headers=headers,
                files={"file": (filename, content)},
                data={"document_version_id": document_version_id},
            )
            response.raise_for_status()
            payload = response.json()
        content_list = payload.get("content_list", [])
        parsed = parse_mineru_content_list(
            document_version_id,
            filename,
            content_list,
            provider_version=str(payload.get("provider_version", "unknown")),
        )
        return parsed.blocks
