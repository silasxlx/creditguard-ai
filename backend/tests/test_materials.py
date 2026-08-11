from __future__ import annotations

from decimal import Decimal
from io import BytesIO

from app.adapters import ParsedBlock
from app.facts import build_fact_payload, detect_conflicts, normalize_value
from app.parsing import (
    parse_docx,
    parse_mineru_content_list,
    parse_pdf,
    parse_xlsx,
)


def _fact_text(term_months: int = 24, amount: int = 1_000_000) -> str:
    return "\n".join(
        [
            "企业名称：合成科技有限公司",
            "统一社会信用代码：91310000SYNTH001",
            "成立日期：2020-01-01",
            "法定代表人：张三",
            "所属行业：制造业",
            f"申请金额：{amount}",
            "币种：CNY",
            f"申请期限：{term_months}个月",
            "贷款用途：流动资金",
            "总资产：5000000",
            "总负债：2000000",
            "流动资产：3000000",
            "流动负债：1500000",
            "营业收入：8000000",
            "净利润：600000",
        ]
    )


def _pdf_bytes(text: str) -> bytes:
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text, fontsize=9)
    return document.tobytes()


def _docx_bytes(text: str) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("授信申请", level=1)
    document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes(text: str) -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    assert sheet is not None
    for line in text.splitlines():
        key, value = line.split("：", 1)
        sheet.append([key, value])
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def test_local_parsers_preserve_evidence_locators() -> None:
    text = _fact_text()
    pdf = parse_pdf("doc-pdf-v1", _pdf_bytes(text), "application.pdf")
    assert pdf.quality.needs_mineru is False
    assert pdf.blocks
    assert pdf.blocks[0].page == 1
    assert pdf.blocks[0].locator["page"] == 1
    assert pdf.blocks[0].evidence_id

    docx = parse_docx("doc-docx-v1", _docx_bytes(text), "application.docx")
    assert docx.blocks[0].section_path == ["授信申请"]
    assert "paragraph_index" in docx.blocks[1].locator

    xlsx = parse_xlsx("doc-xlsx-v1", _xlsx_bytes(text), "financial.xlsx")
    assert xlsx.blocks
    assert xlsx.blocks[0].block_type == "table"
    assert xlsx.blocks[0].locator["sheet"] == "Sheet"
    assert xlsx.quality.cache_missing is False


def test_mineru_content_list_keeps_page_bbox_and_provider_version() -> None:
    parsed = parse_mineru_content_list(
        "doc-mineru-v1",
        "scanned.pdf",
        [{"page_idx": 2, "bbox": [1, 2, 3, 4], "type": "text", "text": "申请金额：1000000"}],
        provider_version="mineru-2.1",
    )
    assert parsed.quality.parser == "mineru"
    assert parsed.blocks[0].page == 3
    assert parsed.blocks[0].locator["bbox"] == [1, 2, 3, 4]
    assert parsed.blocks[0].provider == "mineru:mineru-2.1"


def test_fact_extraction_normalizes_all_standard_fields() -> None:
    from app.parsing import ParsedDocument, ParseQuality

    text = _fact_text()
    parsed = ParsedDocument(
        "doc-facts-v1",
        "application.pdf",
        [
            ParsedBlock(
                document_version_id="doc-facts-v1",
                block_id="block-facts-1",
                evidence_id="evidence-facts-1",
                page=1,
                text=text,
                source="application.pdf",
                locator={"page": 1},
            )
        ],
        ParseQuality("test", len(text), 1, len(text), len(text)),
    )
    payload = build_fact_payload([parsed])
    assert set(payload["fields"]) == {f"F{i:02d}" for i in range(1, 16)}
    assert payload["missing_fields"] == []
    assert payload["requires_review"] is False
    assert payload["fields"]["F08"]["selected_value"] == 24
    assert payload["fields"]["F06"]["selected_value"] == "1000000"


def test_conflict_threshold_is_absolute_and_relative_for_numeric_values() -> None:
    left = ParsedBlock(
        document_version_id="doc-a",
        block_id="block-a",
        evidence_id="evidence-a",
        text="申请金额：1000000",
        source="a.pdf",
        locator={"page": 1},
    )
    right = ParsedBlock(
        document_version_id="doc-b",
        block_id="block-b",
        evidence_id="evidence-b",
        text="申请金额：1020000",
        source="b.pdf",
        locator={"page": 1},
    )
    from app.facts import _candidate_from_block

    conflicts = detect_conflicts(
        [
            _candidate_from_block("F06", left, "1000000"),
            _candidate_from_block("F06", right, "1020000"),
        ]
    )
    assert len(conflicts) == 1
    assert conflicts[0].material is True
    assert conflicts[0].difference is not None
    assert Decimal(conflicts[0].difference["absolute"]) == Decimal("20000")
    assert Decimal(conflicts[0].difference["relative"]) > Decimal("0.01")

    assert normalize_value("F06", "100.5万元") == ("1005000", [])
